"""导出工具模块。

该模块提供了将分析结果导出为不同文件格式（如CSV、Word）的静态方法。
它将导出逻辑与UI层分离，提高了代码的模块化和可维护性。
"""

import io
import pandas as pd
from docx import Document
from docx.shared import Inches
from PIL import Image
import numpy as np
from typing import Dict, List, Any

class Exporter:
    """
    一个包含各种导出功能的静态工具类。
    """
    @staticmethod
    def export_to_csv(details_data: List[Dict[str, Any]], filepath: str) -> None:
        """将详细数据导出为CSV文件。

        Args:
            details_data (List[Dict[str, Any]]): 包含详细测量结果的字典列表。
            filepath (str): 保存CSV文件的路径。
        """
        if not details_data:
            raise ValueError("无法导出空的详细数据。")
        
        details_df = pd.DataFrame(details_data)
        details_df.to_csv(filepath, index=False, encoding='utf-8-sig')

    @staticmethod
    def export_to_word(
        summary_data: Dict[str, Any], 
        image_data: np.ndarray, 
        key_map: Dict[str, str], 
        filepath: str
    ) -> None:
        """将完整的分析结果（摘要、详情、图像）导出为Word文档。

        Args:
            summary_data (Dict[str, Any]): 包含摘要和详细信息的测量结果字典。
            image_data (np.ndarray): 用于可视化的结果图像。
            key_map (Dict[str, str]): 用于将程序键名翻译为用户友好标签的字典。
            filepath (str): 保存Word文档的路径。
        """
        doc = Document()
        doc.add_heading('岩心分析结果报告', level=1)

        # 添加摘要
        doc.add_heading('分析摘要', level=2)
        for key, value in summary_data.items():
            if key == 'details': continue
            label = key_map.get(key, key)
            if isinstance(value, float):
                value_str = f"{value:.4f}"
            else:
                value_str = str(value)
            doc.add_paragraph(f"{label}: {value_str}")

        # 添加详细数据表格
        doc.add_heading('详细数据', level=2)
        details_data = summary_data.get('details', [])
        if details_data:
            headers = list(details_data[0].keys())
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = key_map.get(h, h)

            for item in details_data:
                row_cells = table.add_row().cells
                for i, key in enumerate(headers):
                    value = item.get(key)
                    if isinstance(value, float):
                        value_str = f"{value:.4f}"
                    else:
                        value_str = str(value)
                    row_cells[i].text = value_str

        # 添加可视化图像
        doc.add_heading('可视化结果', level=2)
        if image_data is not None:
            pil_img = Image.fromarray(image_data)
            img_stream = io.BytesIO()
            pil_img.save(img_stream, format='PNG')
            img_stream.seek(0)
            doc.add_picture(img_stream, width=Inches(6.0))

        doc.save(filepath) 